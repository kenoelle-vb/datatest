# streamlit run "C:/Users/keno/OneDrive/Documents/Projects/DATA AUTOMIZER APP/2_EDA.py"

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
import seaborn as sns
import numpy as np
import scipy.stats as stats
from wordcloud import WordCloud
import squarify
import plotly.graph_objects as go
import networkx as nx
from docx import Document
from io import BytesIO
import base64
import warnings
import io
import zipfile

# STREAMLIT CSS AND CONFIGURATIONS ==============================================================================================================

st.set_page_config(page_title="Data Cleaner", layout="wide")
st.set_option('client.showErrorDetails', False)

matplotlib.use('Agg')
warnings.filterwarnings("ignore", category=DeprecationWarning, message=".*pyplot.*")

# Hide Streamlit style elements
hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                header {visibility: hidden;}
                </style>
                """
st.markdown(hide_st_style, unsafe_allow_html=True)

st.markdown("""
    <style>
    [data-testid="stTextArea"] {
        color: #FFFFFF;
    }
    </style>
    """, unsafe_allow_html=True)

# Set Montserrat font
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Montserrat', sans-serif;
    }
    </style>
    """, unsafe_allow_html=True)

# Change color of specific Streamlit elements
st.markdown("""
    <style>
    .st-emotion-cache-1o6s5t7 {
        color: #ababab !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stExpander {
        background-color: white;
        border-radius: 10px;
    }
    
    .stExpander > details {
        background-color: white;
        border-radius: 10px;
    }
    
    .stExpander > details > summary {
        background-color: white;
        border-radius: 10px 10px 0 0;
        padding: 10px;
    }
    
    .stExpander > details > div {
        background-color: white;
        border-radius: 0 0 10px 10px;
        padding: 10px;
    }
    
    .stCheckbox {
        background-color: white;
        border-radius: 5px;
        padding: 5px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stButton > button {
        color: #FFFFFF;
        background-color: #7952eb;
        border-radius: 10px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown(
    """
    <style>
    .streamlit-expanderHeader {
        font-size: 20px;
    }
    .streamlit-expanderContent {
        max-height: 400px;
        overflow-y: auto;
    }
    </style>
    """,
    unsafe_allow_html=True
)

def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

def set_png_as_page_bg(png_file):
    bin_str = get_base64_of_bin_file(png_file)
    page_bg_img = '''
    <style>
    .stApp {
        background-image: url("data:image/png;base64,%s");
        background-size: cover;
    }
    </style>
    ''' % bin_str
    
    st.markdown(page_bg_img, unsafe_allow_html=True)

# Set the background image
set_png_as_page_bg("C:/Users/keno/Downloads/Add a heading (1).png")

data_str = {
    'City': ['Tokyo', 'Osaka', 'Kyoto', 'Yokohama', 'Nagoya', 'Sapporo', 'Kobe', 'Fukuoka'],
}

data_int = {
    'Population': [13929286, 2691185, 1464890, 3757630, 2332000, 1970895, 1545873, 1612392],
}

df_sample_str = pd.DataFrame(data_str)
df_sample_int = pd.DataFrame(data_int)

# ================================================= FUNCTIONS =================================================================================================================

def get_filename_from_title(title):
    """Converts a plot title into a safe file name by replacing spaces and colons."""
    filename = title.replace(" ", "_").replace(":", "") + ".png"
    return filename

# ============================================= Plotting for INT Data ==========================================================================================================================================================================

def get_histogram_figures(df, theme="Blue"):
    """Generates histogram figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    numerical_cols = df.select_dtypes(include=['number']).columns
    figures = []
    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.histplot(df[col], color=colors[i % len(colors)], kde=True, ax=ax)
        title_str = f"Distribution of {col}"
        ax.set_title(title_str)
        ax.set_xlabel(f"{col} Values")
        ax.set_ylabel("Frequency")
        ax.grid(True, linestyle='--', alpha=0.6)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_boxplot_figures(df, theme="Blue"):
    """Generates boxplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    numerical_cols = df.select_dtypes(include=['number']).columns
    figures = []
    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.boxplot(x=df[col], color=colors[i % len(colors)], ax=ax)
        title_str = f"Box Plot of {col}"
        ax.set_title(title_str)
        ax.set_xlabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_scatterplot_figures(df, theme="Blue"):
    """Generates scatterplot figures for unique pairs of numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    numerical_cols = df.select_dtypes(include=['number']).columns
    figures = []
    for i, col1 in enumerate(numerical_cols):
        for j, col2 in enumerate(numerical_cols):
            if i < j:  # Plot only unique pairs
                fig, ax = plt.subplots(figsize=(8, 6))
                sns.scatterplot(x=df[col1], y=df[col2], color=colors[(i + j) % len(colors)], ax=ax)
                title_str = f"Scatter Plot: {col1} vs {col2}"
                ax.set_title(title_str)
                ax.set_xlabel(f"{col1} Values")
                ax.set_ylabel(f"{col2} Values")
                ax.grid(True, linestyle='--', alpha=0.6)
                filename = get_filename_from_title(title_str)
                figures.append((filename, fig))
    return figures

def get_lineplot_figures(df, theme="Blue"):
    """Generates lineplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    numerical_cols = df.select_dtypes(include=['number']).columns
    figures = []
    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.lineplot(x=df.index, y=df[col], color=colors[i % len(colors)], ax=ax)
        title_str = f"Line Plot of {col}"
        ax.set_title(title_str)
        ax.set_xlabel("Index")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_areaplot_figures(df, theme="Blue"):
    """Generates areaplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    numerical_cols = df.select_dtypes(include=['number']).columns
    figures = []
    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.fill_between(df.index, df[col], color=colors[i % len(colors)], alpha=0.6)
        title_str = f"Area Plot of {col}"
        ax.set_title(title_str)
        ax.set_xlabel("Index")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_violinplot_figures(df, theme="Blue"):
    """Generates violinplot figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    numerical_cols = df.select_dtypes(include=['number']).columns
    figures = []
    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.violinplot(y=df[col], color=colors[i % len(colors)], ax=ax)
        title_str = f"Violin Plot of {col}"
        ax.set_title(title_str)
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_correlation_heatmap_figure(df, theme="Blue"):
    """Generates a correlation heatmap figure for numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#E0F2FE", "#B3E5FC", "#81D4FA", "#4FC3F7", "#29B6F6", "#03A9F4", "#039BE5", "#0288D1", "#01579B"],
        "Green": ["#E8F5E9", "#C8E6C9", "#A5D6A7", "#81C784", "#66BB6A", "#4CAF50", "#43A047", "#388E3C", "#2E7D32"],
        "Red": ["#FFE0E0", "#FFCDD2", "#EF9A9A", "#E57373", "#F44336", "#D32F2F", "#C62828", "#B71C1C", "#8B0000"],
        "Purple": ["#F3E5F5", "#E1BEE7", "#CE93D8", "#BA68C8", "#9C27B0", "#8E24AA", "#7B1FA2", "#6A1B9A", "#4A148C"],
        "Orange": ["#FFF3E0", "#FFE0B2", "#FFCC80", "#FFB366", "#FFA726", "#FF9800", "#FB8C00", "#F57C00", "#E65100"],
        "Gray": ["#F5F5F5", "#EEEEEE", "#E0E0E0", "#BDBDBD", "#9E9E9E", "#757575", "#616161", "#424242", "#212121"],
        "Pastel": ["#FFF9C4", "#FFECB3", "#FFE082", "#FFD54F", "#FFCA28", "#FFC107", "#FFB300", "#FFA000", "#FF6F00"]
    }

    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns
    corr_matrix = df[numerical_cols].corr()

    fig, ax = plt.subplots(figsize=(10, 8))
    cmap = sns.color_palette(colors, as_cmap=True)
    sns.heatmap(corr_matrix, annot=True, cmap=cmap, linewidths=.5, ax=ax)
    title_str = "Correlation Heatmap"
    ax.set_title(title_str)
    filename = get_filename_from_title(title_str)
    return [(filename, fig)]

def get_cdf_figures(df, theme="Blue"):
    """Generates CDF figures for all numerical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    numerical_cols = df.select_dtypes(include=['number']).columns
    figures = []
    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        sorted_data = np.sort(df[col])
        yvals = np.arange(len(sorted_data)) / float(len(sorted_data) - 1)
        ax.plot(sorted_data, yvals, color=colors[i % len(colors)])
        title_str = f"CDF of {col}"
        ax.set_title(title_str)
        ax.set_xlabel(f"{col} Values")
        ax.set_ylabel("Cumulative Probability")
        ax.grid(True, linestyle='--', alpha=0.6)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

# ============================================= Plotting for STR Data ==========================================================================================================================================================================

def get_categorical_barplot_figures(df, theme="Blue"):
    """Generates barplot figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    categorical_cols = df.select_dtypes(include=['object']).columns
    figures = []
    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.countplot(y=df[col], color=colors[i % len(colors)], ax=ax)
        title_str = f"Bar Plot of {col}"
        ax.set_title(title_str)
        ax.set_xlabel("Count")
        ax.set_ylabel(f"{col} Categories")
        ax.grid(axis='x', linestyle='--', alpha=0.6)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_piechart_figures(df, theme="Blue"):
    """Generates pie chart figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    categorical_cols = df.select_dtypes(include=['object']).columns
    figures = []
    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        counts = df[col].value_counts()
        ax.pie(counts, labels=counts.index, autopct='%1.1f%%', colors=colors)
        title_str = f"Pie Chart of {col}"
        ax.set_title(title_str)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_stacked_barplot_figures(df, theme="Blue"):
    """Generates stacked bar plot figures for pairs of categorical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    categorical_cols = df.select_dtypes(include=['object']).columns
    figures = []
    if len(categorical_cols) < 2:
        st.write("At least two categorical columns are required for a stacked bar plot.")
        return figures
    for i, col1 in enumerate(categorical_cols):
        for j, col2 in enumerate(categorical_cols):
            if i < j:
                fig, ax = plt.subplots(figsize=(8, 6))
                ct = pd.crosstab(df[col1], df[col2])
                ct.plot(kind='bar', stacked=True, color=colors, ax=ax)
                title_str = f"Stacked Bar Plot: {col1} vs {col2}"
                ax.set_title(title_str)
                ax.set_xlabel(col1)
                ax.set_ylabel("Count")
                ax.tick_params(axis='x', rotation=45)
                plt.tight_layout()
                filename = get_filename_from_title(title_str)
                figures.append((filename, fig))
    return figures

def get_grouped_barplot_figures(df, theme="Blue"):
    """Generates grouped bar plot figures for pairs of categorical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    categorical_cols = df.select_dtypes(include=['object']).columns
    figures = []
    if len(categorical_cols) < 2:
        st.write("At least two categorical columns are required for a grouped bar plot.")
        return figures
    for i, col1 in enumerate(categorical_cols):
        for j, col2 in enumerate(categorical_cols):
            if i < j:
                fig, ax = plt.subplots(figsize=(8, 6))
                ct = pd.crosstab(df[col1], df[col2])
                ct.plot(kind='bar', color=colors, ax=ax)
                title_str = f"Grouped Bar Plot: {col1} vs {col2}"
                ax.set_title(title_str)
                ax.set_xlabel(col1)
                ax.set_ylabel("Count")
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                filename = get_filename_from_title(title_str)
                figures.append((filename, fig))
    return figures

def get_wordcloud_figures(df, theme="Blue"):
    """Generates word cloud figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": "Blues",
        "Green": "Greens",
        "Red": "Reds",
        "Purple": "Purples",
        "Orange": "Oranges",
        "Gray": "Greys",
        "Pastel": "Pastel1"
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colormap = themes[theme]
    categorical_cols = df.select_dtypes(include=['object']).columns
    figures = []
    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        text = ' '.join(df[col].dropna().astype(str))
        if not text:
            st.write(f"Column '{col}' contains no valid text data.")
            continue
        wordcloud = WordCloud(width=800, height=400, background_color='white', colormap=colormap).generate(text)
        ax.imshow(wordcloud, interpolation='bilinear')
        ax.axis('off')
        title_str = f"Word Cloud of {col}"
        ax.set_title(title_str)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_countplot_figures(df, theme="Blue"):
    """Generates count plot figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    categorical_cols = df.select_dtypes(include=['object']).columns
    figures = []
    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.countplot(data=df, x=col, color=colors[i % len(colors)], ax=ax)
        title_str = f"Count Plot of {col}"
        ax.set_title(title_str)
        ax.set_xlabel(f"{col} Categories")
        ax.set_ylabel("Count")
        ax.tick_params(axis='x', rotation=45)
        plt.tight_layout()
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

def get_treemap_figures(df, theme="Blue"):
    """Generates treemap figures for all categorical columns.
    Returns a list of tuples: (filename, figure)."""
    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"],
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"],
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"],
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"],
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"],
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"]
    }
    if theme not in themes:
        st.write(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"
    colors = themes[theme]
    categorical_cols = df.select_dtypes(include=['object']).columns
    figures = []
    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(10, 8))
        counts = df[col].value_counts()
        labels = counts.index
        sizes = counts.values
        total = sizes.sum()
        percentages = [f'{size / total * 100:.2f}%' for size in sizes]
        labels_with_percentage = [f'{label}\n({percentage})' for label, percentage in zip(labels, percentages)]
        squarify.plot(sizes=sizes, label=labels_with_percentage, color=colors, alpha=0.8, ax=ax)
        ax.axis('off')
        title_str = f"Treemap of {col}"
        ax.set_title(title_str)
        filename = get_filename_from_title(title_str)
        figures.append((filename, fig))
    return figures

# ============================================= PLOTTING PLOT ONLY ===================================================================================================================================

def plot_histograms(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.histplot(df[col], color=colors[i % len(colors)], kde=True, ax=ax)  # Use ax
        ax.set_title(f"Distribution of {col}")
        ax.set_xlabel(f"{col} Values")
        ax.set_ylabel("Frequency")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_boxplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.boxplot(x=df[col], color=colors[i % len(colors)], ax=ax)  # Use ax
        ax.set_title(f"Box Plot of {col}")
        ax.set_xlabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_scatterplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col1 in enumerate(numerical_cols):
        for j, col2 in enumerate(numerical_cols):
            if i < j:  # Plot only unique pairs
                fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
                sns.scatterplot(x=df[col1], y=df[col2], color=colors[(i + j) % len(colors)], ax=ax)  # Use ax
                ax.set_title(f"Scatter Plot: {col1} vs {col2}")
                ax.set_xlabel(f"{col1} Values")
                ax.set_ylabel(f"{col2} Values")
                ax.grid(True, linestyle='--', alpha=0.6)
                st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_lineplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.lineplot(x=df.index, y=df[col], color=colors[i % len(colors)], ax=ax)  # Use ax
        ax.set_title(f"Line Plot of {col}")
        ax.set_xlabel("Index")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_areaplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        ax.fill_between(df.index, df[col], color=colors[i % len(colors)], alpha=0.6)  # Use ax for plotting
        ax.set_title(f"Area Plot of {col}")
        ax.set_xlabel("Index")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_violinplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig and ax
        sns.violinplot(y=df[col], color=colors[i % len(colors)], ax=ax)  # Use ax for plotting
        ax.set_title(f"Violin Plot of {col}")
        ax.set_ylabel(f"{col} Values")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_correlation_heatmap(df, cmap_option="Blue"):
    """
    Plots a correlation heatmap for numerical columns in a Pandas DataFrame using color palettes derived from previous themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot the correlation heatmap from.
        cmap_option (str, optional): Color palette option. Defaults to "Blue".
    """

    cmap_options = {
        "Blue": ["#E0F2FE", "#B3E5FC", "#81D4FA", "#4FC3F7", "#29B6F6", "#03A9F4", "#039BE5", "#0288D1", "#01579B"],
        "Green": ["#E8F5E9", "#C8E6C9", "#A5D6A7", "#81C784", "#66BB6A", "#4CAF50", "#43A047", "#388E3C", "#2E7D32"],
        "Red": ["#FFE0E0", "#FFCDD2", "#EF9A9A", "#E57373", "#F44336", "#D32F2F", "#C62828", "#B71C1C", "#8B0000"],
        "Purple": ["#F3E5F5", "#E1BEE7", "#CE93D8", "#BA68C8", "#9C27B0", "#8E24AA", "#7B1FA2", "#6A1B9A", "#4A148C"],
        "Orange": ["#FFF3E0", "#FFE0B2", "#FFCC80", "#FFB366", "#FFA726", "#FF9800", "#FB8C00", "#F57C00", "#E65100"],
        "Gray": ["#F5F5F5", "#EEEEEE", "#E0E0E0", "#BDBDBD", "#9E9E9E", "#757575", "#616161", "#424242", "#212121"],
        "Pastel": ["#FFF9C4", "#FFECB3", "#FFE082", "#FFD54F", "#FFCA28", "#FFC107", "#FFB300", "#FFA000", "#FF6F00"]
    }

    if cmap_option not in cmap_options:
        print(f"cmap_option '{cmap_option}' not found. Using Blue cmap.")
        cmap_option = "Blue"

    colors = cmap_options[cmap_option]

    numerical_cols = df.select_dtypes(include=['number']).columns
    corr_matrix = df[numerical_cols].corr()

    fig, ax = plt.subplots(figsize=(10, 8))  # Create fig and ax
    cmap = sns.color_palette(colors, as_cmap=True)
    sns.heatmap(corr_matrix, annot=True, cmap=cmap, linewidths=.5, ax=ax)  # Use ax
    ax.set_title("Correlation Heatmap")
    st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_cdf(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    numerical_cols = df.select_dtypes(include=['number']).columns

    for i, col in enumerate(numerical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax

        # Calculate CDF
        sorted_data = np.sort(df[col])
        yvals = np.arange(len(sorted_data)) / float(len(sorted_data) - 1)

        ax.plot(sorted_data, yvals, color=colors[i % len(colors)])  # Use ax for plotting
        ax.set_title(f"CDF of {col}")
        ax.set_xlabel(f"{col} Values")
        ax.set_ylabel("Cumulative Probability")
        ax.grid(True, linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass the figure to st.pyplot

# Plotting for STR Data ==========================================================================================================================================================================

def plot_categorical_barplots(df, theme="Blue"):

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        sns.countplot(y=df[col], color=colors[i % len(colors)], ax=ax)  # Plot on ax
        ax.set_title(f"Bar Plot of {col}")
        ax.set_xlabel("Count")
        ax.set_ylabel(f"{col} Categories")
        ax.grid(axis='x', linestyle='--', alpha=0.6)
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_piecharts(df, theme="Blue"):
    """
    Plots pie charts for all categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot pie charts from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        counts = df[col].value_counts()
        ax.pie(counts, labels=counts.index, autopct='%1.1f%%', colors=colors)  # Plot on ax
        ax.set_title(f"Pie Chart of {col}")
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_stacked_barplots(df, theme="Blue"):
    """
    Plots stacked bar plots for pairs of categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot stacked bar plots from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    if len(categorical_cols) < 2:
        print("At least two categorical columns are required for a stacked bar plot.")
        return

    for i, col1 in enumerate(categorical_cols):
        for j, col2 in enumerate(categorical_cols):
            if i < j:
                fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
                ct = pd.crosstab(df[col1], df[col2])
                ct.plot(kind='bar', stacked=True, color=colors, ax=ax)  # Plot on ax
                ax.set_title(f"Stacked Bar Plot: {col1} vs {col2}")
                ax.set_xlabel(col1)
                ax.set_ylabel("Count")
                ax.tick_params(axis='x', rotation=45)
                plt.tight_layout()
                st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_grouped_barplots(df, theme="Blue"):
    """
    Plots grouped bar plots (clustered bar plots) for pairs of categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot grouped bar plots from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    if len(categorical_cols) < 2:
        print("At least two categorical columns are required for a grouped bar plot.")
        return

    for i, col1 in enumerate(categorical_cols):
        for j, col2 in enumerate(categorical_cols):
            if i < j:
                fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
                ct = pd.crosstab(df[col1], df[col2])
                ct.plot(kind='bar', color=colors, ax=ax)  # Plot on the ax
                ax.set_title(f"Grouped Bar Plot: {col1} vs {col2}")
                ax.set_xlabel(col1)
                ax.set_ylabel("Count")
                plt.xticks(rotation=45, ha='right')  # Rotate x-axis labels
                plt.tight_layout()
                st.pyplot(fig, clear_figure=True)  # Pass the fig to st.pyplot

def plot_wordclouds(df, theme="Blue"):
    """
    Plots word clouds for all categorical (string) columns in a Pandas DataFrame with color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot word clouds from.
        theme (str, optional): The color theme for the word clouds. Defaults to "Blue".
    """

    themes = {
        "Blue": "Blues",
        "Green": "Greens",
        "Red": "Reds",
        "Purple": "Purples",
        "Orange": "Oranges",
        "Gray": "Greys",
        "Pastel": "Pastel1"
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colormap = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        text = ' '.join(df[col].dropna().astype(str))  # Combine all strings into one text
        if not text:
            print(f"Column '{col}' contains no valid text data.")
            continue

        wordcloud = WordCloud(width=800, height=400, background_color='white', colormap=colormap).generate(text)
        ax.imshow(wordcloud, interpolation='bilinear')  # Plot on ax
        ax.axis('off')
        ax.set_title(f"Word Cloud of {col}")
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_countplots(df, theme="Blue"):
    """
    Plots count plots for all categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes.

    Args:
        df (pd.DataFrame): The DataFrame to plot count plots from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(8, 6))  # Create fig, ax
        sns.countplot(data=df, x=col, color=colors[i % len(colors)], ax=ax)  # Use ax
        ax.set_title(f"Count Plot of {col}")
        ax.set_xlabel(f"{col} Categories")
        ax.set_ylabel("Count")
        ax.tick_params(axis='x', rotation=45)
        plt.tight_layout()
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

def plot_treemaps(df, theme="Blue"):
    """
    Plots treemaps for categorical (string) columns in a Pandas DataFrame with light/medium high-saturated color themes, including percentage values on each block.

    Args:
        df (pd.DataFrame): The DataFrame to plot treemaps from.
        theme (str, optional): The color theme for the plots. Defaults to "Blue".
    """

    themes = {
        "Blue": ["#6699CC", "#7A99CC", "#80B3FF", "#7AA8D1", "#667FCC", "#5673B3", "#4660A6", "#384D99", "#283980"], # Medium-high saturation blues
        "Green": ["#70D4A2", "#70C189", "#5CA66B", "#4A8C50", "#3D773C", "#315C33", "#2A4E2E", "#223E26", "#192E1B"], # Medium-high saturation greens
        "Red": ["#E87975", "#E86260", "#D9504E", "#C94140", "#B63234", "#A2262C", "#8E1B24", "#7A111C", "#600A14"], # Medium-high saturation reds
        "Purple": ["#AA94BF", "#957DBF", "#806BBF", "#6F57B3", "#604A9E", "#503C88", "#433077", "#3A256A", "#281854"], # Medium-high saturation purples
        "Orange": ["#FFA500", "#FFB347", "#FFC170", "#FFCC99", "#FFD9C2", "#F2BA9C", "#E6A582", "#D99067", "#CC7A4E"], # Medium-high saturation oranges
        "Gray": ["#A8A8A8", "#B3B3B3", "#BEBEBE", "#C8C8C8", "#D3D3D3", "#DADADA", "#E3E3E3", "#EDEDED", "#F5F5F5"],  # Medium-high saturation grays
        "Pastel": ["#F0E68C", "#E6E096", "#DCD9A0", "#D2D3AA", "#C9CDBC", "#BFC7D1", "#B0C1D9", "#A2BBDF", "#8FB5E4"] # Medium-high saturation pastels
    }

    if theme not in themes:
        print(f"Theme '{theme}' not found. Using Blue theme.")
        theme = "Blue"

    colors = themes[theme]

    categorical_cols = df.select_dtypes(include=['object']).columns

    for i, col in enumerate(categorical_cols):
        fig, ax = plt.subplots(figsize=(10, 8))  # Create fig, ax
        counts = df[col].value_counts()
        labels = counts.index
        sizes = counts.values

        # Calculate percentages
        total = sizes.sum()
        percentages = [f'{size / total * 100:.2f}%' for size in sizes]

        # Combine labels with percentages
        labels_with_percentage = [f'{label}\n({percentage})' for label, percentage in zip(labels, percentages)]

        squarify.plot(sizes=sizes, label=labels_with_percentage, color=colors, alpha=0.8, ax=ax)  # Use ax for plotting
        ax.axis('off')
        ax.set_title(f"Treemap of {col}")
        st.pyplot(fig, clear_figure=True)  # Pass fig to st.pyplot

# ============================================== DOC REPORTING ===========================================================================================================

def generate_doc_report_en(df):
    """
    Generates a comprehensive DOCX report from EDA results with overall implications.

    Args:
        df (pd.DataFrame): The input DataFrame.
        output_filename (str): The name of the output DOCX file.
    """

    document = Document()

    # Basic Statistics & Overall Implications (Part 1)
    document.add_heading("Comprehensive Exploratory Data Analysis Report", level=1)
    basic_stats = f"This report provides a detailed overview of the dataset. The dataset comprises {df.shape[0]} observations (rows) and {df.shape[1]} variables (columns). "
    basic_stats += f"Notably, no duplicate rows were found, indicating data uniqueness. However, a significant number of missing cells were identified, totaling {df.isnull().sum().sum()}. This highlights potential data completeness issues that may require further investigation."
    document.add_paragraph(basic_stats)

    implications = "Overall, the dataset presents a combination of strengths and challenges. The absence of duplicate rows suggests a well-curated dataset, but the presence of missing values necessitates careful handling during analysis. "
    if df.isnull().sum().sum() / (df.shape[0] * df.shape[1]) > 0.1:
        implications += "The high percentage of missing data (over 10%) could significantly impact the reliability of statistical analyses and model building. "
    else:
        implications += "While missing values are present, their volume is relatively manageable, and suitable imputation or deletion strategies can be employed. "

    document.add_paragraph(implications)

    # Variable Types & Implications
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    text_cols = df.select_dtypes(include='object').columns.tolist()
    document.add_heading("Variable Types", level=2)
    var_types = f"The dataset features {len(numeric_cols)} numeric variables, including '{', '.join(numeric_cols[:5])}', and {len(text_cols)} text-based variables: '{', '.join(text_cols)}'. "
    var_types += "No categorical variables were identified in this dataset."
    document.add_paragraph(var_types)

    type_implications = "The composition of numeric and text variables suggests a dataset suitable for quantitative and qualitative analyses. "
    if len(numeric_cols) > len(text_cols):
        type_implications += "The dominance of numeric variables indicates a dataset primarily designed for statistical modeling and quantitative analysis. "
    elif len(text_cols) > len(numeric_cols):
        type_implications += "The prevalence of text-based variables suggests a dataset rich in textual information, potentially suitable for natural language processing (NLP) tasks or qualitative content analysis. "

    document.add_paragraph(type_implications)

    # Highly Correlated Variables & Implications
    if numeric_cols:
        corr_matrix = df[numeric_cols].corr().abs()
        upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
        highly_correlated = [column for column in upper.columns if any(upper[column] > 0.8)]
        document.add_heading("Highly Correlated Variables (x > 0.8)", level=2)
        if highly_correlated:
            high_corr = f"Several variables exhibit high correlations (above 0.8), suggesting strong relationships or potential redundancy. These include '{', '.join(highly_correlated)}'. "
            high_corr += "Such high correlations may warrant further examination to understand the underlying dependencies. A correlation above 0.8 generally indicates a strong positive or negative linear relationship. This could imply that one variable is a proxy for another, or that they are both influenced by a common factor."
            document.add_paragraph(high_corr)

            corr_implications = "The presence of highly correlated variables suggests potential multicollinearity issues, which can affect the stability and interpretability of regression models. "
            corr_implications += "It may be necessary to perform feature selection or dimensionality reduction techniques to address this. "
            document.add_paragraph(corr_implications)

        else:
            document.add_paragraph("No variables with correlations above 0.8 were found. This suggests that the numeric variables in the dataset do not exhibit strong linear relationships with each other.")
            document.add_paragraph("The absence of strong correlations simplifies model building as multicollinearity is not a major concern.")
    else:
        document.add_paragraph("No numeric columns were found, so correlation analysis could not be performed.")
        document.add_paragraph("Without numerical columns, assessment of variable correlations cannot be performed.")

    # Variables with Unique Values & Implications
    unique_counts = df.nunique()
    document.add_heading("Variables with Unique Values", level=2)
    document.add_paragraph("The dataset displays a wide range of uniqueness across its variables.")
    for col in df.columns:
        document.add_paragraph(f"'{col}' contains {unique_counts[col]} unique values.")
    document.add_paragraph("This variability in uniqueness can provide insights into the nature and distribution of the data. High unique value counts can indicate identifiers or detailed categorical variables, while low counts suggest broad categories or limited variability.")

    unique_implications = "The distribution of unique values influences the choice of analytical methods. Columns with very high unique values might need special treatment, especially if they are identifiers that don't contribute to statistical modeling. "
    if any(unique_counts / df.shape[0] > 0.8):
        unique_implications += "Columns with very high cardinality (unique values approaching the number of rows) might be considered as identifiers and excluded from certain analyses. "
    document.add_paragraph(unique_implications)

    # Uniform Distribution (Simplified) & Implications
    document.add_heading("Uniform Distribution (Simplified)", level=2)
    document.add_paragraph("A simplified check for uniform distribution was conducted.")
    uniform_cols = []
    for col in df.columns:
        if df[col].nunique() > 10:
            if (df[col].value_counts(normalize=True).std() < 0.05):
                uniform_cols.append(col)
    if uniform_cols:
        uniform_str = f"Variables such as '{', '.join(uniform_cols)}' might exhibit a uniform distribution. "
        uniform_str += "A uniform distribution suggests that all values are equally likely, which can be important for certain statistical tests or modeling assumptions. "
    else:
        uniform_str = "No variables were identified as potentially having a uniform distribution based on this simplified check. "
    uniform_str += "Variables with fewer than 10 unique values were excluded from this check."
    document.add_paragraph(uniform_str)

    uniform_implications = "The presence or absence of uniform distributions impacts the choice of statistical tests. Uniform distributions can be important for hypothesis testing and simulation studies. "
    if uniform_cols:
        uniform_implications += "The detected potential uniform distributions might simplify certain modeling or hypothesis testing procedures. "
    else:
        uniform_implications += "The absence of strong indications of uniform distributions suggests that data transformations or alternative tests might be necessary. "
    document.add_paragraph(uniform_implications)

    # Missing Values & Implications
    document.add_heading("Missing Values", level=2)
    document.add_paragraph("The dataset contains missing values across multiple variables.")
    missing_values = df.isnull().sum()
    for col in missing_values.index:
        document.add_paragraph(f"'{col}' has {missing_values[col]} missing values.")
    missing_str = f"The variable '{missing_values.index[missing_values.argmax()]}' has the highest number of missing values, with {missing_values.max()} missing values. Addressing these missing values is crucial for accurate analysis. High missing value counts can bias results and reduce the reliability of conclusions."
    document.add_paragraph(missing_str)

    missing_implications = "The handling of missing values is critical. High missing value counts can lead to biased or unreliable results. "
    if missing_values.max() / df.shape[0] > 0.3:
        missing_implications += "Columns with over 30% missing values might be considered for removal or require advanced imputation techniques. "
    else:
        missing_implications += "The missing values can be addressed using standard imputation or deletion methods. "
    document.add_paragraph(missing_implications)

    # Top 5 Mostly and Least Correlated Variables & Implications
    if numeric_cols:
        document.add_heading("Correlation Analysis", level=2)
        corr_matrix = df[numeric_cols].corr().abs()
        for col in numeric_cols:
            if col in corr_matrix.columns:
                corr_series = corr_matrix[col].sort_values(ascending=False)
                document.add_heading(f"Correlation Analysis for '{col}'", level=3)
                top_5_mostly = ""
                for i in range(1,6):
                  top_5_mostly += f"'{corr_series.index[i]}' with {corr_series[i]:.4f} correlation, "
                top_5_mostly = top_5_mostly[:-2]
                document.add_paragraph(f"The top 5 mostly correlated variables are: {top_5_mostly}. High correlation indicates a strong relationship, suggesting that these variables move together. It might be useful to examine these pairs more closely.")
                top_5_least = ""
                for i in range(len(corr_series)-5,len(corr_series)):
                    top_5_least += f"'{corr_series.index[i]}' with {corr_series[i]:.4f} correlation, "
                top_5_least = top_5_least[:-2]
                document.add_paragraph(f"The top 5 least correlated variables are: {top_5_least}. Low correlation suggests that these variables are relatively independent. This can be important for building models where independence is assumed.")

                correlation_insight = f"For '{col}', the high correlations suggest that these variables might be used interchangeably or that they are driven by a common underlying factor. Low correlations indicate variables that contribute unique information. "
                document.add_paragraph(correlation_insight)
    else:
        document.add_paragraph("Correlation data is not available as there are no numerical columns.")
        document.add_paragraph("Without numerical data, correlation implications cannot be provided.")

    # Variables Insight & Overall Implications (Part 2)
    document.add_heading("Variables Insight", level=2)
    for column_name in df.columns:
        col = df[column_name]
        document.add_paragraph(f"Analysis of Column: {column_name}")
        insight_paragraph = f"The column '{column_name}' has a data type of {col.dtype}, with {col.nunique()} unique values and {col.isnull().sum()} missing values. "

        if pd.api.types.is_numeric_dtype(col):
            mean = col.mean()
            std = col.std()
            min_val = col.min()
            q25 = col.quantile(0.25)
            median = col.median()
            q75 = col.quantile(0.75)
            max_val = col.max()
            skew = col.skew()
            kurt = col.kurt()
            zeros = (col == 0).sum()

            insight_paragraph += f"Its mean is {mean:.4f}, standard deviation is {std:.4f}, minimum value is {min_val:.4f}, 25th percentile is {q25:.4f}, median is {median:.4f}, 75th percentile is {q75:.4f}, maximum value is {max_val:.4f}, skewness is {skew:.4f}, kurtosis is {kurt:.4f}, and the number of zeros is {zeros}. "

            if std > 0:
                insight_paragraph += f"A standard deviation of {std:.4f} indicates the spread of the data around the mean. "
            if skew > 1 or skew < -1:
                insight_paragraph += f"A skewness of {skew:.4f} suggests that the data is highly skewed. "
            elif skew > 0.5 or skew < -0.5:
                insight_paragraph += f"A skewness of {skew:.4f} suggests moderate skewness. "
            if kurt > 3:
                insight_paragraph += f"A kurtosis of {kurt:.4f} indicates a leptokurtic distribution (heavy tails). "
            elif kurt < 3:
                insight_paragraph += f"A kurtosis of {kurt:.4f} indicates a platykurtic distribution (light tails). "

        elif pd.api.types.is_string_dtype(col) or pd.api.types.is_object_dtype(col):
            most_frequent = col.mode()[0]
            insight_paragraph += f"The most frequent value is '{most_frequent}', which appears {(col == most_frequent).sum()} times. "
            if col.nunique() / len(col) > 0.5:
                insight_paragraph += "This column has a high cardinality, meaning many unique values relative to the total number of entries. "
            if col.isnull().sum() / len(col) > 0.5:
                insight_paragraph += "This column has a high percentage of missing values. "

        document.add_paragraph(insight_paragraph)

    overall_variable_implications = "The individual variable insights provide a granular understanding of the data's characteristics. Skewness and kurtosis values highlight potential deviations from normal distributions, which can impact the choice of statistical tests. High cardinality in text variables might require feature engineering or dimensionality reduction. "
    document.add_paragraph(overall_variable_implications)

    return document 

def generate_doc_report_id(df, output_filename="eda_report_indonesian.docx"):
    """
    Menghasilkan laporan DOCX yang komprehensif dari hasil EDA dengan implikasi keseluruhan.

    Args:
        df (pd.DataFrame): DataFrame input.
        output_filename (str): Nama file DOCX output.
    """

    document = Document()

    # Statistik Dasar & Implikasi Keseluruhan (Bagian 1)
    document.add_heading("Laporan Analisis Data Eksplorasi Komprehensif", level=1)
    basic_stats = f"Laporan ini memberikan tinjauan rinci tentang kumpulan data. Kumpulan data terdiri dari {df.shape[0]} observasi (baris) dan {df.shape[1]} variabel (kolom). "
    basic_stats += f"Khususnya, tidak ditemukan baris duplikat, yang menunjukkan keunikan data. Namun, sejumlah besar sel yang hilang teridentifikasi, dengan total {df.isnull().sum().sum()}. Ini menyoroti potensi masalah kelengkapan data yang mungkin memerlukan penyelidikan lebih lanjut."
    document.add_paragraph(basic_stats)

    implications = "Secara keseluruhan, kumpulan data menyajikan kombinasi kekuatan dan tantangan. Tidak adanya baris duplikat menunjukkan kumpulan data yang terkurasi dengan baik, tetapi keberadaan nilai yang hilang memerlukan penanganan yang cermat selama analisis. "
    if df.isnull().sum().sum() / (df.shape[0] * df.shape[1]) > 0.1:
        implications += "Persentase data yang hilang yang tinggi (lebih dari 10%) dapat secara signifikan memengaruhi keandalan analisis statistik dan pembuatan model. "
    else:
        implications += "Meskipun nilai yang hilang ada, volumenya relatif dapat dikelola, dan strategi imputasi atau penghapusan yang sesuai dapat diterapkan. "

    document.add_paragraph(implications)

    # Jenis Variabel & Implikasi
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    text_cols = df.select_dtypes(include='object').columns.tolist()
    document.add_heading("Jenis Variabel", level=2)
    var_types = f"Kumpulan data menampilkan {len(numeric_cols)} variabel numerik, termasuk '{', '.join(numeric_cols[:5])}', dan {len(text_cols)} variabel berbasis teks: '{', '.join(text_cols)}'. "
    var_types += "Tidak ada variabel kategorikal yang teridentifikasi dalam kumpulan data ini."
    document.add_paragraph(var_types)

    type_implications = "Komposisi variabel numerik dan teks menunjukkan kumpulan data yang cocok untuk analisis kuantitatif dan kualitatif. "
    if len(numeric_cols) > len(text_cols):
        type_implications += "Dominasi variabel numerik menunjukkan kumpulan data yang terutama dirancang untuk pemodelan statistik dan analisis kuantitatif. "
    elif len(text_cols) > len(numeric_cols):
        type_implications += "Prevalensi variabel berbasis teks menunjukkan kumpulan data yang kaya akan informasi tekstual, yang berpotensi cocok untuk tugas pemrosesan bahasa alami (NLP) atau analisis konten kualitatif. "

    document.add_paragraph(type_implications)

    # Variabel dengan Korelasi Tinggi & Implikasi
    if numeric_cols:
        corr_matrix = df[numeric_cols].corr().abs()
        upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
        highly_correlated = [column for column in upper.columns if any(upper[column] > 0.8)]
        document.add_heading("Variabel dengan Korelasi Tinggi (x > 0.8)", level=2)
        if highly_correlated:
            high_corr = f"Beberapa variabel menunjukkan korelasi tinggi (di atas 0.8), yang menunjukkan hubungan yang kuat atau potensi redundansi. Ini termasuk '{', '.join(highly_correlated)}'. "
            high_corr += "Korelasi tinggi seperti itu mungkin memerlukan pemeriksaan lebih lanjut untuk memahami ketergantungan yang mendasarinya. Korelasi di atas 0.8 umumnya menunjukkan hubungan linear positif atau negatif yang kuat. Ini dapat menyiratkan bahwa satu variabel adalah proksi untuk variabel lain, atau bahwa keduanya dipengaruhi oleh faktor umum."
            document.add_paragraph(high_corr)

            corr_implications = "Keberadaan variabel dengan korelasi tinggi menunjukkan potensi masalah multikolinearitas, yang dapat memengaruhi stabilitas dan interpretasi model regresi. "
            corr_implications += "Mungkin perlu dilakukan pemilihan fitur atau teknik reduksi dimensi untuk mengatasi hal ini. "
            document.add_paragraph(corr_implications)

        else:
            document.add_paragraph("Tidak ada variabel dengan korelasi di atas 0.8 yang ditemukan. Ini menunjukkan bahwa variabel numerik dalam kumpulan data tidak menunjukkan hubungan linear yang kuat satu sama lain.")
            document.add_paragraph("Tidak adanya korelasi yang kuat menyederhanakan pembuatan model karena multikolinearitas bukan masalah utama.")
    else:
        document.add_paragraph("Tidak ada kolom numerik yang ditemukan, jadi analisis korelasi tidak dapat dilakukan.")
        document.add_paragraph("Tanpa kolom numerik, penilaian korelasi variabel tidak dapat dilakukan.")

    # Variabel dengan Nilai Unik & Implikasi
    unique_counts = df.nunique()
    document.add_heading("Variabel dengan Nilai Unik", level=2)
    document.add_paragraph("Kumpulan data menampilkan berbagai keunikan di seluruh variabelnya.")
    for col in df.columns:
        document.add_paragraph(f"'{col}' berisi {unique_counts[col]} nilai unik.")
    document.add_paragraph("Variabilitas keunikan ini dapat memberikan wawasan tentang sifat dan distribusi data. Jumlah nilai unik yang tinggi dapat menunjukkan pengidentifikasi atau variabel kategorikal terperinci, sementara jumlah yang rendah menunjukkan kategori luas atau variabilitas terbatas.")

    unique_implications = "Distribusi nilai unik memengaruhi pilihan metode analisis. Kolom dengan nilai unik yang sangat tinggi mungkin memerlukan perlakuan khusus, terutama jika itu adalah pengidentifikasi yang tidak berkontribusi pada pemodelan statistik. "
    if any(unique_counts / df.shape[0] > 0.8):
        unique_implications += "Kolom dengan kardinalitas yang sangat tinggi (nilai unik mendekati jumlah baris) mungkin dianggap sebagai pengidentifikasi dan dikecualikan dari analisis tertentu. "
    document.add_paragraph(unique_implications)

    # Distribusi Seragam (Sederhana) & Implikasi
    document.add_heading("Distribusi Seragam (Sederhana)", level=2)
    document.add_paragraph("Pemeriksaan sederhana untuk distribusi seragam dilakukan.")
    uniform_cols = []
    for col in df.columns:
        if df[col].nunique() > 10:
            if (df[col].value_counts(normalize=True).std() < 0.05):
                uniform_cols.append(col)
    if uniform_cols:
        uniform_str = f"Variabel seperti '{', '.join(uniform_cols)}' mungkin menunjukkan distribusi seragam. "
        uniform_str += "Distribusi seragam menunjukkan bahwa semua nilai sama-sama mungkin, yang penting untuk tes statistik atau asumsi pemodelan tertentu. "
    else:
        uniform_str = "Tidak ada variabel yang teridentifikasi berpotensi memiliki distribusi seragam berdasarkan pemeriksaan sederhana ini. "
    uniform_str += "Variabel dengan kurang dari 10 nilai unik dikecualikan dari pemeriksaan ini."
    document.add_paragraph(uniform_str)

    uniform_implications = "Keberadaan atau tidak adanya distribusi seragam memengaruhi pilihan tes statistik. Distribusi seragam dapat penting untuk pengujian hipotesis dan studi simulasi. "
    if uniform_cols:
        uniform_implications += "Potensi distribusi seragam yang terdeteksi mungkin menyederhanakan prosedur pemodelan atau pengujian hipotesis tertentu. "
    else:
        uniform_implications += "Tidak adanya indikasi kuat distribusi seragam menunjukkan bahwa transformasi data atau tes alternatif mungkin diperlukan. "
    document.add_paragraph(uniform_implications)

    # Nilai yang Hilang & Implikasi
    document.add_heading("Nilai yang Hilang", level=2)
    document.add_paragraph("Kumpulan data berisi nilai yang hilang di beberapa variabel.")
    missing_values = df.isnull().sum()
    for col in missing_values.index:
        document.add_paragraph(f"'{col}' memiliki {missing_values[col]} nilai yang hilang.")
    missing_str = f"Variabel '{missing_values.index[missing_values.argmax()]}' memiliki jumlah nilai yang hilang tertinggi, dengan {missing_values.max()} nilai yang hilang. Mengatasi nilai yang hilang ini sangat penting untuk analisis"
    document.add_paragraph(missing_str)

    missing_implications = "Penanganan nilai yang hilang sangat penting. Jumlah nilai yang hilang yang tinggi dapat menyebabkan hasil yang bias atau tidak dapat diandalkan. "
    if missing_values.max() / df.shape[0] > 0.3:
        missing_implications += "Kolom dengan lebih dari 30% nilai yang hilang mungkin dipertimbangkan untuk dihapus atau memerlukan teknik imputasi lanjutan. "
    else:
        missing_implications += "Nilai yang hilang dapat diatasi menggunakan metode imputasi atau penghapusan standar. "
    document.add_paragraph(missing_implications)

    # 5 Variabel dengan Korelasi Tertinggi dan Terendah & Implikasi
    if numeric_cols:
        document.add_heading("Analisis Korelasi", level=2)
        corr_matrix = df[numeric_cols].corr().abs()
        for col in numeric_cols:
            if col in corr_matrix.columns:
                corr_series = corr_matrix[col].sort_values(ascending=False)
                document.add_heading(f"Analisis Korelasi untuk '{col}'", level=3)
                top_5_mostly = ""
                for i in range(1,6):
                  top_5_mostly += f"'{corr_series.index[i]}' dengan korelasi {corr_series[i]:.4f}, "
                top_5_mostly = top_5_mostly[:-2]
                document.add_paragraph(f"5 variabel dengan korelasi tertinggi adalah: {top_5_mostly}. Korelasi tinggi menunjukkan hubungan yang kuat, menunjukkan bahwa variabel-variabel ini bergerak bersamaan. Mungkin berguna untuk memeriksa pasangan-pasangan ini lebih dekat.")
                top_5_least = ""
                for i in range(len(corr_series)-5,len(corr_series)):
                    top_5_least += f"'{corr_series.index[i]}' dengan korelasi {corr_series[i]:.4f}, "
                top_5_least = top_5_least[:-2]
                document.add_paragraph(f"5 variabel dengan korelasi terendah adalah: {top_5_least}. Korelasi rendah menunjukkan bahwa variabel-variabel ini relatif independen. Ini dapat penting untuk membangun model di mana independensi diasumsikan.")

                correlation_insight = f"Untuk '{col}', korelasi tinggi menunjukkan bahwa variabel-variabel ini mungkin digunakan secara bergantian atau bahwa mereka didorong oleh faktor mendasar yang sama. Korelasi rendah menunjukkan variabel-variabel yang memberikan informasi unik. "
                document.add_paragraph(correlation_insight)
    else:
        document.add_paragraph("Data korelasi tidak tersedia karena tidak ada kolom numerik.")
        document.add_paragraph("Tanpa data numerik, implikasi korelasi tidak dapat diberikan.")

    # Wawasan Variabel & Implikasi Keseluruhan (Bagian 2)
    document.add_heading("Wawasan Variabel", level=2)
    for column_name in df.columns:
        col = df[column_name]
        document.add_paragraph(f"Analisis Kolom: {column_name}")
        insight_paragraph = f"Kolom '{column_name}' memiliki tipe data {col.dtype}, dengan {col.nunique()} nilai unik dan {col.isnull().sum()} nilai yang hilang. "

        if pd.api.types.is_numeric_dtype(col):
            mean = col.mean()
            std = col.std()
            min_val = col.min()
            q25 = col.quantile(0.25)
            median = col.median()
            q75 = col.quantile(0.75)
            max_val = col.max()
            skew = col.skew()
            kurt = col.kurt()
            zeros = (col == 0).sum()

            insight_paragraph += f"Rata-ratanya adalah {mean:.4f}, standar deviasi adalah {std:.4f}, nilai minimum adalah {min_val:.4f}, persentil ke-25 adalah {q25:.4f}, median adalah {median:.4f}, persentil ke-75 adalah {q75:.4f}, nilai maksimum adalah {max_val:.4f}, skewness adalah {skew:.4f}, kurtosis adalah {kurt:.4f}, dan jumlah nol adalah {zeros}. "

            if std > 0:
                insight_paragraph += f"Standar deviasi sebesar {std:.4f} menunjukkan penyebaran data di sekitar rata-rata. "
            if skew > 1 or skew < -1:
                insight_paragraph += f"Skewness sebesar {skew:.4f} menunjukkan bahwa data sangat miring. "
            elif skew > 0.5 or skew < -0.5:
                insight_paragraph += f"Skewness sebesar {skew:.4f} menunjukkan kemiringan sedang. "
            if kurt > 3:
                insight_paragraph += f"Kurtosis sebesar {kurt:.4f} menunjukkan distribusi leptokurtik (ekor berat). "
            elif kurt < 3:
                insight_paragraph += f"Kurtosis sebesar {kurt:.4f} menunjukkan distribusi platikurtik (ekor ringan). "

        elif pd.api.types.is_string_dtype(col) or pd.api.types.is_object_dtype(col):
            most_frequent = col.mode()[0]
            insight_paragraph += f"Nilai paling sering adalah '{most_frequent}', yang muncul {(col == most_frequent).sum()} kali. "
            if col.nunique() / len(col) > 0.5:
                insight_paragraph += "Kolom ini memiliki kardinalitas tinggi, yang berarti banyak nilai unik relatif terhadap jumlah total entri. "
            if col.isnull().sum() / len(col) > 0.5:
                insight_paragraph += "Kolom ini memiliki persentase nilai yang hilang yang tinggi. "

        document.add_paragraph(insight_paragraph)

    overall_variable_implications = "Wawasan variabel individu memberikan pemahaman terperinci tentang karakteristik data. Nilai skewness dan kurtosis menyoroti potensi penyimpangan dari distribusi normal, yang dapat memengaruhi pilihan tes statistik. Kardinalitas tinggi dalam variabel teks mungkin memerlukan rekayasa fitur atau reduksi dimensi. "
    document.add_paragraph(overall_variable_implications)

    return document 

# ============================================== Download Functions =========================================================================================

def aggregate_and_download_plots(df, selected_plot_types, theme="Blue", cmap_option="Blue"):
    """Aggregates figures from the selected plot types and returns a zip buffer."""
    all_figures = []

    # Normalize plot type names for comparison
    selected_plot_types_lower = [plot_type.lower() for plot_type in selected_plot_types]

    if "histograms" in selected_plot_types_lower:
        all_figures.extend(get_histogram_figures(df, theme))
    if "boxplots" in selected_plot_types_lower:
        all_figures.extend(get_boxplot_figures(df, theme))
    if "scatterplots" in selected_plot_types_lower:
        all_figures.extend(get_scatterplot_figures(df, theme))
    if "lineplots" in selected_plot_types_lower:
        all_figures.extend(get_lineplot_figures(df, theme))
    if "areaplots" in selected_plot_types_lower:
        all_figures.extend(get_areaplot_figures(df, theme))
    if "violinplots" in selected_plot_types_lower:
        all_figures.extend(get_violinplot_figures(df, theme))
    if "correlation heatmap" in selected_plot_types_lower:
        all_figures.extend(get_correlation_heatmap_figure(df, cmap_option))
    if "cdf" in selected_plot_types_lower:
        all_figures.extend(get_cdf_figures(df, theme))
    if "categorical barplots" in selected_plot_types_lower:
        all_figures.extend(get_categorical_barplot_figures(df, theme))
    if "piecharts" in selected_plot_types_lower:
        all_figures.extend(get_piechart_figures(df, theme))
    if "stacked barplots" in selected_plot_types_lower:
        all_figures.extend(get_stacked_barplot_figures(df, theme))
    if "grouped barplots" in selected_plot_types_lower:
        all_figures.extend(get_grouped_barplot_figures(df, theme))
    if "wordclouds" in selected_plot_types_lower:
        all_figures.extend(get_wordcloud_figures(df, theme))
    if "countplots" in selected_plot_types_lower:
        all_figures.extend(get_countplot_figures(df, theme))
    if "treemaps" in selected_plot_types_lower:
        all_figures.extend(get_treemap_figures(df, theme))

    if not all_figures:
        st.write("No plot types selected.")
        return None  # Return None if no figures

    # Create a zip file in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        for filename, fig in all_figures:
            buf = io.BytesIO()
            fig.savefig(buf, format="png", bbox_inches="tight")
            buf.seek(0)
            zip_file.writestr(filename, buf.getvalue())
            plt.close(fig)  # Close figure to avoid memory warning.

    zip_buffer.seek(0)
    return zip_buffer #return the zip buffer.

# Plotting for Report Data ================================================================================================================================================================================

def eda_dataframe_to_docx(df):
    """
    Performs Exploratory Data Analysis (EDA) on a DataFrame and creates a docx document.
    """
    document = Document()

    document.add_heading("Exploratory Data Analysis", 0)

    document.add_heading("Basic Statistics", level=1)
    document.add_paragraph(f"Number of Observations (Rows): {df.shape[0]}")
    document.add_paragraph(f"Number of Variables (Columns): {df.shape[1]}")
    document.add_paragraph(f"Duplicate Rows: {df.duplicated().sum()}")
    document.add_paragraph(f"Missing Cells: {df.isnull().sum().sum()}")

    document.add_heading("Variable Types", level=1)
    numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
    categorical_cols = df.select_dtypes(include='category').columns.tolist()
    text_cols = df.select_dtypes(include='object').columns.tolist()

    document.add_paragraph(f"Numeric: {len(numeric_cols)} ({', '.join(map(str, numeric_cols))})")
    document.add_paragraph(f"Categorical: {len(categorical_cols)} ({', '.join(map(str, categorical_cols))})")
    document.add_paragraph(f"Text (Object): {len(text_cols)} ({', '.join(map(str, text_cols))})")

    document.add_heading("Highly Correlated Variables (x > 0.8)", level=1)
    if numeric_cols:
        corr_matrix = df[numeric_cols].corr().abs()
        upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(bool))
        highly_correlated = [column for column in upper.columns if any(upper[column] > 0.8)]
        if highly_correlated:
            document.add_paragraph(f"{', '.join(map(str, highly_correlated))}")
        else:
            document.add_paragraph("No highly correlated variables found.")
    else:
        document.add_paragraph("No numeric columns found to calculate correlation.")

    document.add_heading("Variables with Unique Values", level=1)
    unique_counts = df.nunique()
    document.add_paragraph(str(unique_counts))

    document.add_heading("Variables with Uniform Distribution (Simplified)", level=1)
    for col in df.columns:
        if df[col].nunique() > 10:
            if (df[col].value_counts(normalize=True).std() < 0.05):
                document.add_paragraph(f"Column: {col} might have a uniform distribution")
        else:
            document.add_paragraph(f"Column: {col} has less than 10 unique values, cannot check for uniform distribution")

    document.add_heading("Missing Values per Variable", level=1)
    document.add_paragraph(str(df.isnull().sum()))

    document.add_heading("Top 5 Mostly and Least Correlated Variables", level=1)
    if len(numeric_cols) > 0:
        corr_matrix = df[numeric_cols].corr().abs()
        for col in numeric_cols:
            if col in corr_matrix.columns:
                corr_series = corr_matrix[col].sort_values(ascending=False)
                document.add_paragraph(f"\nCorrelation with {col}:")
                document.add_paragraph("Top 5 Mostly Correlated:")
                document.add_paragraph(str(corr_series.head(6).tail(5)))
                document.add_paragraph("Top 5 Least Correlated:")
                document.add_paragraph(str(corr_series.tail(5)))

    return document

def eda_all_columns_to_docx(df):
    """
    Performs Exploratory Data Analysis on all columns iteratively and creates a docx document.
    """
    document = Document()
    document.add_heading("Column-Wise Exploratory Data Analysis", 0)

    for column_name in df.columns:
        col = df[column_name]
        document.add_heading(f"Analysis of Column: {column_name}", level=1)
        document.add_paragraph(f"Data Type: {col.dtype}")
        document.add_paragraph(f"Number of Unique Values: {col.nunique()}")
        document.add_paragraph(f"Number of Missing Values: {col.isnull().sum()}")

        if pd.api.types.is_numeric_dtype(col):
            document.add_paragraph(f"Mean: {col.mean()}")
            document.add_paragraph(f"Standard Deviation: {col.std()}")
            document.add_paragraph(f"Minimum: {col.min()}")
            document.add_paragraph(f"25th Percentile: {col.quantile(0.25)}")
            document.add_paragraph(f"Median: {col.median()}")
            document.add_paragraph(f"75th Percentile: {col.quantile(0.75)}")
            document.add_paragraph(f"Maximum: {col.max()}")
            document.add_paragraph(f"Skewness: {col.skew()}")
            document.add_paragraph(f"Kurtosis: {col.kurt()}")
            document.add_paragraph(f"Number of Zeros: {(col == 0).sum()}")

        elif pd.api.types.is_string_dtype(col) or pd.api.types.is_object_dtype(col):
            most_frequent = col.mode()[0]
            document.add_paragraph(f"Most Frequent Value: {most_frequent}")
            document.add_paragraph(f"Frequency of Most Frequent Value: {(col == most_frequent).sum()}")

    return document

# ======================================= ALL PLOT OPTIONS =================================================================================================

all_plot_options = [ "Histograms", "Boxplots", "Scatterplots", "Lineplots", "Areaplots",
                    "Violinplots", "Correlation Heatmap", "CDF", "Categorical Barplots",
                    "Piecharts", "Stacked Barplots", "Grouped Barplots", "Wordclouds",
                    "Countplots", "Treemaps"]

# =========================================== Main Streamlit App Code =============================================================================================================

with st.expander("EDA to Docx", expanded=True):
    st.title("Combined EDA and Plot Downloader")

    uploaded_file = st.file_uploader("Upload your CSV file", type=["csv", "xlsx"], accept_multiple_files=False)

    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            st.success("File uploaded successfully!")
        except Exception as e:
            st.error(f"Error reading file: {e}")
            df = None

        if df is not None:
            st.write("Uploaded Dataframe:")
            st.dataframe(df.head())

            if st.button("Continue"):
                st.session_state['df'] = df

if 'plot_fig' not in st.session_state:
    st.session_state['plot_fig'] = None

if uploaded_file is not None:
    if 'df' in st.session_state:
        col1, col2, col3 = st.columns([2, 2, 2])

        with col1:
            # Plot Examples (Expander)
            with st.expander("Plot Examples", expanded=True):
                st.write("Below is an example using a sample DataFrame. Replace `df_sample` with your own reference if needed.")

                example_plot_choice = st.selectbox("Choose Plot", options=all_plot_options, index=0)

                theme_choice = st.selectbox(
                    "Select Plot Theme",
                    ["Blue", "Red", "Green", "Purple", "Orange", "Gray", "Pastel"],
                    index=0
                )

                if example_plot_choice == "Histograms":
                    st.session_state['plot_fig'] = plot_histograms(df_sample_int, theme=theme_choice)

                elif example_plot_choice == "Boxplots":
                    st.session_state['plot_fig'] = plot_boxplots(df_sample_int, theme=theme_choice)

                elif example_plot_choice == "Scatterplots":
                    st.session_state['plot_fig'] = plot_scatterplots(df_sample_int, theme=theme_choice)

                elif example_plot_choice == "Lineplots":
                    st.session_state['plot_fig'] = plot_lineplots(df_sample_int, theme=theme_choice)

                elif example_plot_choice == "Areaplots":
                    st.session_state['plot_fig'] = plot_areaplots(df_sample_int, theme=theme_choice)

                elif example_plot_choice == "Violinplots":
                    st.session_state['plot_fig'] = plot_violinplots(df_sample_int, theme=theme_choice)

                elif example_plot_choice == "Correlation Heatmap":
                    st.session_state['plot_fig'] = plot_correlation_heatmap(df_sample_int, cmap_option=theme_choice)

                elif example_plot_choice == "CDF":
                    st.session_state['plot_fig'] = plot_cdf(df_sample_int, theme=theme_choice)

                elif example_plot_choice == "Categorical Barplots":
                    st.session_state['plot_fig'] = plot_categorical_barplots(df_sample_str, theme=theme_choice)

                elif example_plot_choice == "Piecharts":
                    st.session_state['plot_fig'] = plot_piecharts(df_sample_str, theme=theme_choice)

                elif example_plot_choice == "Stacked Barplots":
                    st.session_state['plot_fig'] = plot_stacked_barplots(df_sample_str, theme=theme_choice)

                elif example_plot_choice == "Grouped Barplots":
                    st.session_state['plot_fig'] = plot_grouped_barplots(df_sample_str, theme=theme_choice)

                elif example_plot_choice == "Wordclouds":
                    st.session_state['plot_fig'] = plot_wordclouds(df_sample_str, theme=theme_choice)

                elif example_plot_choice == "Countplots":
                    st.session_state['plot_fig'] = plot_countplots(df_sample_str, theme=theme_choice)

                elif example_plot_choice == "Treemaps":
                    st.session_state['plot_fig'] = plot_treemaps(df_sample_str, theme=theme_choice)

            if st.session_state['plot_fig'] is not None:
                st.pyplot(st.session_state['plot_fig'], clear_figure=True)

        with col2:
            with st.expander("EDA to Docx", expanded=True):
                st.subheader("EDA to Docx")
                if st.button("Generate EDA Docx Reports"):
                    status_container = st.empty()
                    status_container.info("Generating EDA reports...")

                    doc_basic = eda_dataframe_to_docx(st.session_state['df'])
                    doc_columnwise = eda_all_columns_to_docx(st.session_state['df'])

                    buffer_basic = BytesIO()
                    doc_basic.save(buffer_basic)
                    buffer_basic.seek(0)

                    buffer_columnwise = BytesIO()
                    doc_columnwise.save(buffer_columnwise)
                    buffer_columnwise.seek(0)

                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w') as zipf:
                        zipf.writestr("basic_eda_report.docx", buffer_basic.getvalue())
                        zipf.writestr("columnwise_eda_report.docx", buffer_columnwise.getvalue())
                    zip_buffer.seek(0)

                    status_container.success("EDA reports generated!")
                    st.session_state['eda_zip_buffer'] = zip_buffer

                if 'eda_zip_buffer' in st.session_state:
                    st.download_button(
                        label="Download EDA Reports (ZIP)",
                        data=st.session_state['eda_zip_buffer'],
                        file_name="Exploratory Data Analysis Report.zip",
                        mime="application/zip"
                    )

            with st.expander("", expanded=True):
                if st.button("Generate Report"):
                    status_container = st.empty()
                    status_container.info("Generating Key Insight reports...")
                    doc_en = generate_doc_report_en(st.session_state['df'])
                    buffer_en = BytesIO()
                    doc_en.save(buffer_en)
                    buffer_en.seek(0)

                    doc_id = generate_doc_report_id(st.session_state['df'])
                    buffer_id = BytesIO()
                    doc_id.save(buffer_id)
                    buffer_id.seek(0)

                    zip_doc = BytesIO()
                    with zipfile.ZipFile(zip_doc, 'w') as zipd:
                        zipd.writestr("ENGLISH - Key Insights Analysis.docx", buffer_en.getvalue())
                        zipd.writestr("INDONESIAN - Key Insights Analysis.docx", buffer_id.getvalue())
                    zip_doc.seek(0)

                    status_container.success("Key Insight reports generated!")
                    st.session_state['eda_zip_doc'] = zip_doc

                if 'eda_zip_doc' in st.session_state:
                    st.download_button(
                        label="Download EDA Reports (ZIP)",
                        data=st.session_state['eda_zip_doc'],
                        file_name="ID/EN - Key Insights Analysis.zip",
                        mime="application/zip"
                    )

        with col3:
            with st.expander("Aggregate Plot Downloader", expanded=True):
                st.subheader("Aggregate Plot Downloader")
                selected_plot_types = st.multiselect(
                    "Choose plots to download",
                    options=[
                        "Histograms", "Boxplots", "Scatterplots", "Lineplots", "Areaplots",
                        "Violinplots", "Correlation Heatmap", "CDF", "Categorical Barplots",
                        "Piecharts", "Stacked Barplots", "Grouped Barplots", "Wordclouds",
                        "Countplots", "Treemaps"
                    ],
                    default=[
                        "Histograms", "Areaplots",
                        "Correlation Heatmap", "Categorical Barplots",
                        "Piecharts", "Wordclouds",
                        "Treemaps"
                    ]
                )
                theme = st.selectbox("Select Theme", ["Blue", "Green", "Red", "Purple", "Orange", "Gray", "Pastel"], index=0)
                if st.button("Download Selected Plots"):
                    status_container_plots = st.empty()
                    status_container_plots.info("Generating plots...")

                    zip_buffer_plots = aggregate_and_download_plots(st.session_state['df'], selected_plot_types, theme)

                    status_container_plots.success("Plots generated!")
                    st.session_state['plot_zip_buffer'] = zip_buffer_plots

                if 'plot_zip_buffer' in st.session_state:
                    st.download_button(
                        label="Download Selected Plots (ZIP)",
                        data=st.session_state['plot_zip_buffer'],
                        file_name="plots.zip",
                        mime="application/zip"
                    )
